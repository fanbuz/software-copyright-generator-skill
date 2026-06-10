#!/usr/bin/env python3
"""Review a software copyright DOCX package."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

try:
    from docx import Document
except Exception:  # pragma: no cover - exercised only when optional dependency is absent.
    Document = None  # type: ignore[assignment]


def docx_body_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def docx_header_footer_text(doc: Document) -> str:
    parts = []
    for section in doc.sections:
        for area in [section.header, section.footer]:
            parts.extend(p.text for p in area.paragraphs)
            for table in area.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
    return "\n".join(parts)


def count_term(text: str, term: str) -> int:
    if term.isascii() and term.isalnum():
        return len(re.findall(rf"(?<![A-Za-z0-9_]){re.escape(term)}(?![A-Za-z0-9_])", text))
    return text.count(term)


def inspect_txt(path: Path, terms: list[str]) -> dict:
    body = path.read_text(encoding="utf-8", errors="replace")
    combined = "\n".join([path.name, body])
    return {
        "path": str(path),
        "exists": path.exists(),
        "size": path.stat().st_size,
        "paragraphs": sum(1 for line in body.splitlines() if line.strip()),
        "tables": 0,
        "rows_first_table": 0,
        "cols_first_row": 0,
        "images": 0,
        "term_counts": {term: count_term(combined, term) for term in terms},
        "body_term_counts": {term: count_term(body, term) for term in terms},
        "header_footer_term_counts": {term: 0 for term in terms},
    }


def inspect_docx(path: Path, terms: list[str]) -> dict:
    if Document is None:
        return {
            "path": str(path),
            "exists": path.exists(),
            "size": path.stat().st_size if path.exists() else 0,
            "paragraphs": 0,
            "tables": 0,
            "rows_first_table": 0,
            "cols_first_row": 0,
            "images": 0,
            "term_counts": {term: count_term(path.name, term) for term in terms},
            "body_term_counts": {term: 0 for term in terms},
            "header_footer_term_counts": {term: 0 for term in terms},
            "warning": "python-docx unavailable; only filename-level checks were performed",
        }
    doc = Document(str(path))
    body = docx_body_text(doc)
    header_footer = docx_header_footer_text(doc)
    combined = "\n".join([path.name, body, header_footer])
    return {
        "path": str(path),
        "exists": path.exists(),
        "size": path.stat().st_size,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "rows_first_table": len(doc.tables[0].rows) if doc.tables else 0,
        "cols_first_row": len(doc.tables[0].rows[0].cells) if doc.tables and doc.tables[0].rows else 0,
        "images": sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype),
        "term_counts": {term: count_term(combined, term) for term in terms},
        "body_term_counts": {term: count_term(body, term) for term in terms},
        "header_footer_term_counts": {term: count_term(header_footer, term) for term in terms},
    }


def collect_files(directory: Path, patterns: list[str], version: str) -> list[Path]:
    """Match by keyword globs, then require the version anywhere in the filename."""
    found: set[Path] = set()
    for pattern in patterns:
        found.update(directory.glob(pattern))
    return sorted(p for p in found if version in p.name)


def review_package(directory: Path, software_name: str, version: str, forbidden: list[str]) -> dict:
    files = {
        "info": collect_files(
            directory,
            ["*信息采集表*.docx", "*申请表信息*.docx", "*信息采集表*.txt", "*申请表信息*.txt"],
            version,
        ),
        "manual": collect_files(directory, ["*用户手册*.docx", "*操作手册*.docx"], version),
        "source": collect_files(directory, ["*源代码*.docx", "*代码*.docx"], version),
    }
    terms = [software_name, version, *forbidden]
    result = {"dir": str(directory), "files": {}, "missing": [], "warnings": []}
    for kind, matches in files.items():
        if not matches:
            result["missing"].append(kind)
            continue
        # 同类材料优先复核 DOCX；信息件仅有 TXT 时按文本复核并提示。
        docx_matches = [p for p in matches if p.suffix.lower() == ".docx"]
        target = docx_matches[0] if docx_matches else matches[0]
        if target.suffix.lower() == ".txt":
            info = inspect_txt(target, terms)
            result["warnings"].append(f"{kind}: 仅找到 TXT 文本材料；正式信息采集表 DOCX 需通过模板套打生成")
        else:
            info = inspect_docx(target, terms)
        result["files"][kind] = info
        if info.get("warning"):
            result["warnings"].append(f"{kind}: {info['warning']}")
        for forbidden_term in forbidden:
            if info["term_counts"].get(forbidden_term, 0) > 0:
                result["warnings"].append(f"{kind}: forbidden term remains: {forbidden_term}")
        for identity in [software_name, version]:
            if info["term_counts"].get(identity, 0) > 0 and (
                info["body_term_counts"].get(identity, 0) + info["header_footer_term_counts"].get(identity, 0) == 0
            ):
                result["warnings"].append(f"{kind}: {identity} only appears in filename, not body/header/footer")

    result["ok"] = not result["missing"] and all(
        item["term_counts"].get(software_name, 0) > 0 and item["term_counts"].get(version, 0) > 0
        for item in result["files"].values()
    ) and not any("forbidden term remains" in item for item in result["warnings"])
    result["strict_ok"] = result["ok"] and not result["warnings"]
    return result


def write_human_summary(result: dict) -> str:
    lines = ["# 三件套复核摘要", ""]
    lines.append(f"- 目录：{result['dir']}")
    lines.append(f"- ok：{str(result['ok']).lower()}")
    lines.append(f"- strict_ok：{str(result['strict_ok']).lower()}")
    lines.extend(["", "## 文件状态", ""])
    for kind in ("info", "manual", "source"):
        if kind in result["missing"]:
            lines.append(f"- {kind}: 缺失")
        else:
            item = result["files"].get(kind) or {}
            lines.append(f"- {kind}: {item.get('path', '')} ({item.get('size', 0)} bytes)")
    lines.extend(["", "## 风险", ""])
    if result["warnings"]:
        lines.extend(f"- {warning}" for warning in result["warnings"])
    else:
        lines.append("- 无")
    return "\n".join(lines) + "\n"


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--dir", required=True, type=Path)
    parser.add_argument("--software-name", required=True)
    parser.add_argument("--version", required=True)
    parser.add_argument("--forbidden", action="append", default=[])
    parser.add_argument("--summary", type=Path, help="Optional Markdown summary path")
    args = parser.parse_args()

    result = review_package(args.dir, args.software_name, args.version, args.forbidden)
    if args.summary:
        args.summary.parent.mkdir(parents=True, exist_ok=True)
        args.summary.write_text(write_human_summary(result), encoding="utf-8")
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
