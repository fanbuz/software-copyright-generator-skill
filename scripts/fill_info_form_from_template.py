#!/usr/bin/env python3
"""Fill a software-copyright info form by copying and editing a DOCX template."""

from __future__ import annotations

import argparse
import json
import re
import shutil
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import _Cell


def physical_cell(table, row_index: int, cell_index: int) -> _Cell:
    tc = table.rows[row_index]._tr.tc_lst[cell_index]
    return _Cell(tc, table)


def set_cell_text(cell: _Cell, text: str, size: float = 10.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_checkbox_states(row, states: list[bool]) -> None:
    boxes = list(row._tr.iter(qn("w:checkBox")))
    for box, checked in zip(boxes, states):
        default = box.find(qn("w:default"))
        if default is None:
            default = OxmlElement("w:default")
            box.append(default)
        default.set(qn("w:val"), "1" if checked else "0")

        checked_node = box.find(qn("w:checked"))
        if checked_node is None:
            checked_node = OxmlElement("w:checked")
            box.append(checked_node)
        checked_node.set(qn("w:val"), "1" if checked else "0")


def parse_row_map(raw: str | None) -> dict[int, str]:
    if not raw:
        return {}
    data = json.loads(raw)
    return {int(k): str(v) for k, v in data.items()}


def default_values(fields: dict[str, Any]) -> dict[int, str]:
    return {
        1: str(fields.get("software_name", "")),
        2: str(fields.get("short_name", "无")),
        3: str(fields.get("classification", "30100，0000")),
        4: str(fields.get("version", "V1.0")),
        6: str(fields.get("completed_date", "")),
        13: str(fields.get("hardware", "")),
        14: str(fields.get("software_environment", "")),
        15: str(fields.get("languages", "")),
        16: str(fields.get("source_lines", "")),
        17: str(fields.get("main_functions", "")),
        19: str(fields.get("copyright_owner", "")),
        20: str(fields.get("region", "")),
        21: str(fields.get("credit_code", "")),
        24: str(fields.get("applicant", fields.get("copyright_owner", ""))),
        25: str(fields.get("address", "")),
        26: str(fields.get("postcode", "")),
        27: str(fields.get("contact", "")),
        28: str(fields.get("phone", "")),
        29: str(fields.get("email", "")),
        30: str(fields.get("mobile", fields.get("phone", ""))),
        31: str(fields.get("fax", "")),
    }


def fill_form(template: Path, output: Path, fields: dict[str, Any], row_map: dict[int, str]) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, output)
    doc = Document(str(output))
    table = doc.tables[0]

    values = default_values(fields)
    values.update(row_map)
    for row_index, text in values.items():
        if row_index < len(table.rows) and len(table.rows[row_index]._tr.tc_lst) >= 2:
            set_cell_text(physical_cell(table, row_index, 1), text)

    # Common Chinese software-copyright template rows.
    if len(table.rows) >= 10:
        if len(table.rows[8]._tr.tc_lst) >= 3:
            set_cell_text(physical_cell(table, 8, 2), str(fields.get("first_publish_date", "")))
        if len(table.rows[9]._tr.tc_lst) >= 3:
            set_cell_text(physical_cell(table, 9, 2), str(fields.get("first_publish_city", "")))

    checkbox_defaults = {
        "work_description": [True, False],
        "publish_state": [False, True],
        "development_mode": [True, False, False, False],
        "rights_obtain": [True, False, False, False, False, False],
        "enterprise_type": [False, False, True, False, False, False],
    }
    checkbox_rows = {
        "work_description": 5,
        "publish_state": 7,
        "development_mode": 10,
        "rights_obtain": 11,
        "enterprise_type": 22,
    }
    for key, row_index in checkbox_rows.items():
        states = fields.get(f"{key}_checkboxes", checkbox_defaults[key])
        if row_index < len(table.rows):
            set_checkbox_states(table.rows[row_index], [bool(item) for item in states])

    doc.save(str(output))


def structure_report(doc: Document, fields: dict[str, Any], leaks: list[str], output: Path) -> dict[str, Any]:
    table_shapes = []
    for table in doc.tables:
        row_count = len(table.rows)
        first_row_cols = len(table.rows[0].cells) if table.rows else 0
        table_shapes.append({"rows": row_count, "cols_first_row": first_row_cols})
    return {
        "output": str(output),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "table_shapes": table_shapes,
        "rows": len(doc.tables[0].rows) if doc.tables else 0,
        "images": sum(1 for r in doc.part.rels.values() if "image" in r.reltype),
        "blank_preserved_fields": [
            key
            for key in ("completed_date", "first_publish_date", "first_publish_city")
            if key in fields and str(fields.get(key) or "") == ""
        ],
        "forbidden_leaks": leaks,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--fields-json", required=True, type=Path)
    parser.add_argument(
        "--row-map-json",
        help="Optional JSON object mapping zero-based row indexes to replacement text.",
    )
    args = parser.parse_args()

    fields = json.loads(args.fields_json.read_text(encoding="utf-8"))
    fill_form(args.template, args.output, fields, parse_row_map(args.row_map_json))
    doc = Document(str(args.output))
    text = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    old_names = fields.get("forbidden_terms", [])
    leaks = [term for term in old_names if term and re.search(re.escape(str(term)), text)]
    print(json.dumps(structure_report(doc, fields, leaks, args.output), ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
